#!/usr/bin/env python3
"""将docx中的表格转换为md格式，插入到对应章节文件中"""
from docx import Document
import os, re

def table_to_md(table):
    md = []
    header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
    md.append('| ' + ' | '.join(header) + ' |')
    md.append('|' + '|'.join(['---' for _ in header]) + '|')
    for row in table.rows[1:]:
        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        md.append('| ' + ' | '.join(row_data) + ' |')
    return '\n'.join(md)

docx_path = '/Users/liangquan/Desktop/server/githubagentwork/repos/ch-bak/FitPal/lw/0504.docx'
output_dir = '/Users/liangquan/Desktop/server/githubagentwork/repos/ch-bak/FitPal/docs/paper2'

doc = Document(docx_path)

# 获取文档body中所有元素的顺序
body = doc.element.body
elements = []
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        elements.append(('para', child))
    elif tag == 'tbl':
        elements.append(('table', child))

# 建立段落索引映射
para_map = {}
for i, p in enumerate(doc.paragraphs):
    para_map[id(p._element)] = i

# 建立表格索引映射
table_map = {}
for i, t in enumerate(doc.tables):
    table_map[id(t._element)] = i

# 找到Heading 1位置
h1_indices = []
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1':
        h1_indices.append(i)

chapter_names = []
for idx in h1_indices:
    name = doc.paragraphs[idx].text.strip().replace('/', '-').replace('\\', '-')
    chapter_names.append(name)

# 按章节提取内容（含表格，按文档顺序）
for i, (start_idx, chapter_name) in enumerate(zip(h1_indices, chapter_names)):
    end_idx = h1_indices[i+1] if i+1 < len(h1_indices) else len(doc.paragraphs)
    
    content = []
    
    for elem_type, elem in elements:
        if elem_type == 'para':
            p_id = id(elem)
            if p_id in para_map:
                p_idx = para_map[p_id]
                if start_idx <= p_idx < end_idx:
                    content.append(doc.paragraphs[p_idx].text)
        elif elem_type == 'table':
            t_id = id(elem)
            if t_id in table_map:
                # 表格位置估算：找到前一个段落的索引
                prev_para_idx = -1
                for pe_type, pe_elem in elements:
                    if pe_elem is elem:
                        break
                    if pe_type == 'para' and id(pe_elem) in para_map:
                        prev_para_idx = para_map[id(pe_elem)]
                
                if start_idx <= prev_para_idx < end_idx:
                    content.append(table_to_md(doc.tables[table_map[t_id]]))
    
    filename = f'{i+1}-{chapter_name}.md'
    filepath = os.path.join(output_dir, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(content))
    
    print(f'已更新: {filename}')

print('全部完成')
