#!/usr/bin/env python3
"""
将 Markdown 文件转换为 Word 文档 (docx)
表格使用三线表格式，宽度占满，单元格宽度按内容自适应
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import re
import sys

def set_chinese_font(run, font_name='宋体', font_size=10.5):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)

def set_heading_font(run, font_name='黑体', font_size=14):
    """设置标题字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = True

def set_table_header_font(run, font_name='黑体', font_size=10.5):
    """设置表头字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = True

def set_cell_shading(cell, color='D9D9D9'):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def create_three_line_table(doc, headers, rows, table_name=''):
    """
    创建三线表
    三线表格式：顶线、表头底线、底线三条横线，无竖线
    """
    # 创建表格
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表格宽度占满页面
    table.allow_autofit = False
    table.autofit = False
    
    # 获取页面宽度（A4约为15cm可用宽度）
    page_width = Cm(15)
    
    # 计算每列的基础宽度（平均分配）
    col_width = page_width / len(headers)
    
    # 设置列宽
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_width
    
    # 填充表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header.strip()
        # 设置表头字体
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_table_header_font(run)
    
    # 填充数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row_cells):
                cell = row_cells[col_idx]
                cell.text = cell_text.strip()
                # 设置单元格字体和格式
                for paragraph in cell.paragraphs:
                    # 第一列左对齐，其他居中
                    if col_idx == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_chinese_font(run)
    
    # 应用三线表格式（通过边框设置）
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # 设置表格边框为三线表样式
    # 移除所有竖线，只保留三条横线
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # 创建单元格边框
            tcBorders = OxmlElement('w:tcBorders')
            
            # 顶线：第一行或数据行顶部
            top = OxmlElement('w:top')
            if row_idx == 0:  # 表格顶线（最粗）
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')
                top.set(qn('w:color'), '000000')
            elif row_idx == 1:  # 表头底线（较粗）
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '8')
                top.set(qn('w:color'), '000000')
            else:
                top.set(qn('w:val'), 'nil')
            tcBorders.append(top)
            
            # 底线：最后一行
            bottom = OxmlElement('w:bottom')
            if row_idx == len(table.rows) - 1:  # 表格底线（较粗）
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')
            else:
                bottom.set(qn('w:val'), 'nil')
            tcBorders.append(bottom)
            
            # 左边框（无）
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')
            tcBorders.append(left)
            
            # 右边框（无）
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')
            tcBorders.append(right)
            
            tcPr.append(tcBorders)
    
    # 表头行添加底边框（三线表中间那条线）
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
    
    # 在表格后添加空行
    doc.add_paragraph()
    
    return table

def parse_markdown_table(lines, start_idx):
    """解析 Markdown 表格，返回表头、数据行和结束索引"""
    headers = []
    rows = []
    
    # 第一行是表头
    if start_idx < len(lines):
        header_line = lines[start_idx].strip()
        if header_line.startswith('|'):
            headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
    
    # 第二行是分隔符，跳过
    current_idx = start_idx + 2
    
    # 读取数据行直到遇到空行或非表格行
    while current_idx < len(lines):
        line = lines[current_idx].strip()
        if not line or not line.startswith('|'):
            break
        
        # 解析数据行
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        current_idx += 1
    
    return headers, rows, current_idx

def parse_bold_text(text):
    """解析 Markdown 粗体标记 **text**"""
    parts = []
    pattern = r'\*\*(.*?)\*\*'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # 添加普通文本
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        # 添加粗体文本
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    # 添加剩余文本
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    if not parts:
        parts.append(('normal', text))
    
    return parts

def md_to_docx(md_file, docx_file):
    """将 Markdown 文件转换为 Word 文档"""
    
    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 处理一级标题 (# 标题)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_heading_font(run, '黑体', 16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 处理二级标题 (## 标题)
        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_heading_font(run, '黑体', 14)
            i += 1
            continue
        
        # 处理三级标题 (### 标题)
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_heading_font(run, '黑体', 12)
            i += 1
            continue
        
        # 处理表格 (| 开头)
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, next_idx = parse_markdown_table(lines, i)
            if headers:
                create_three_line_table(doc, headers, rows)
            i = next_idx
            continue
        
        # 处理普通段落文本
        # 处理 **粗体** 标记
        p = doc.add_paragraph()
        parts = parse_bold_text(line)
        
        for style, text in parts:
            run = p.add_run(text)
            set_chinese_font(run, '宋体', 12)
            if style == 'bold':
                run.bold = True
        
        # 处理换行（段内换行用<br>表示）
        while i + 1 < len(lines) and lines[i + 1].strip().startswith('<br>'):
            i += 1
            p.add_run().add_break()
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"已生成文档: {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        md_file = '/Users/liangquan/Desktop/server/githubagentwork/repos/ch-bak/FitPal/docs/paper/03-系统分析-v2.md'
        docx_file = '/Users/liangquan/Desktop/server/githubagentwork/repos/ch-bak/FitPal/docs/paper/03-系统分析-v2.docx'
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[1].replace('.md', '.docx')
    
    md_to_docx(md_file, docx_file)
