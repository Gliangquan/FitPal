#!/usr/bin/env python3
"""将 04-4-2.md 的 4.2 系统功能模块设计 转为 docx，插入流程图图片"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

PUML_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'puml')
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '04-系统功能模块设计.docx')

# 图片映射：图编号 -> PNG文件名
IMG_MAP = {
    '图4-2': 'flow-04-02.png',
    '图4-3': 'flow-04-03.png',
    '图4-4': 'flow-04-04.png',
    '图4-5': 'flow-04-05.png',
    '图4-6': 'flow-04-06.png',
    '图4-7': 'flow-04-07.png',
    '图4-8': 'flow-04-08.png',
    '图4-9': 'flow-04-09.png',
    '图4-10': 'flow-04-10.png',
    '图4-11': 'flow-04-11.png',
    '图4-12': 'flow-04-12.png',
    '图4-13': 'flow-04-13.png',
    '图4-14': 'flow-04-14.png',
    '图4-15': 'flow-04-15.png',
    '图4-16': 'flow-04-16.png',
    '图4-17': 'flow-04-17.png',
    '图4-18': 'flow-04-18.png',
    '图4-19': 'flow-04-19.png',
    '图4-20': 'flow-04-20.png',
    '图4-21': 'flow-04-21.png',
    '图4-22': 'flow-04-22.png',
    '图4-23': 'flow-04-23.png',
}


def set_run_font(run, name='宋体', size=Pt(12), bold=False):
    run.font.size = size
    run.font.name = name
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_heading_styled(doc, text, level):
    """添加标题并设置字体"""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 2:
        set_run_font(run, '黑体', Pt(16), bold=True)
    elif level == 3:
        set_run_font(run, '黑体', Pt(14), bold=True)
    elif level == 4:
        set_run_font(run, '黑体', Pt(12), bold=True)
    return h


def add_body_text(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, '宋体', Pt(12))
    return p


def add_image_with_caption(doc, img_path, caption):
    """插入图片 + 图片标题"""
    # 图片段落
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        run = p_img.add_run()
        run.add_picture(img_path, width=Cm(12))
    else:
        run = p_img.add_run(f'[图片缺失: {os.path.basename(img_path)}]')
        set_run_font(run, '宋体', Pt(10))

    # 标题段落
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(caption)
    set_run_font(run, '宋体', Pt(10.5))
    return p_cap


def is_figure_line(line):
    """判断是否为图片引用行"""
    for key in IMG_MAP:
        if line.startswith(key):
            return key, line
    return None, None


def build_docx():
    doc = Document()

    # 读取源文件
    md_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '04-4-2.md')
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = [l.rstrip() for l in f.readlines()]

    for line in lines:
        if not line.strip():
            continue

        # 4.2 标题
        if line.startswith('4.2 '):
            add_heading_styled(doc, line, 2)
            continue

        # 4.2.x 标题
        if line.startswith('4.2.'):
            add_heading_styled(doc, line, 3)
            continue

        # 数字编号子标题（如 "1. 登录注册模块"）
        stripped = line.strip()
        if len(stripped) > 2 and stripped[0].isdigit() and '. ' in stripped[:5]:
            parts = stripped.split('. ', 1)
            if '模块' in parts[1]:
                add_heading_styled(doc, stripped, 4)
                continue

        # 图片行
        fig_key, fig_caption = is_figure_line(line)
        if fig_key:
            img_path = os.path.join(PUML_DIR, IMG_MAP[fig_key])
            add_image_with_caption(doc, img_path, fig_caption)
            continue

        # 正文
        add_body_text(doc, line)

    doc.save(OUTPUT)
    print(f'已生成: {OUTPUT}')


if __name__ == '__main__':
    build_docx()
